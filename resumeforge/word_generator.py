"""Gerador de currículo Word (.docx) a partir de dados estruturados."""

import gc
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import OUTPUT_DIR, SECTION_LABELS
from .models import ResumeData


def generate_word(resume: ResumeData, output_name: str = "resume_tailored") -> Path:
    """Gera um arquivo .docx a partir dos dados do currículo."""
    doc = Document()
    
    # Configurações de estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Margens menores para caber em 1 página se possível
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume.personal.name)
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_details = []
    if resume.personal.phone:
        contact_details.append(resume.personal.phone)
    if resume.personal.email:
        contact_details.append(resume.personal.email)
    if resume.personal.linkedin:
        contact_details.append(resume.personal.linkedin)
    if resume.personal.github:
        contact_details.append(resume.personal.github)
    if resume.personal.location:
        contact_details.append(resume.personal.location)
    
    if contact_details:
        contact_para = doc.add_paragraph(" | ".join(contact_details))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linha divisória
    doc.add_paragraph().add_run("_" * 80).bold = True

    # Resumo
    if resume.summary:
        heading = doc.add_heading(SECTION_LABELS["summary"], level=1)
        doc.add_paragraph(resume.summary)

    # Experiência
    if resume.experience:
        doc.add_heading(SECTION_LABELS["experience"], level=1)
        for exp in resume.experience:
            p = doc.add_paragraph()
            p.add_run(f"{exp.role}").bold = True
            p.add_run(f" | {exp.company}")
            if exp.location:
                p.add_run(f" - {exp.location}")
            
            p_date = doc.add_paragraph()
            p_date.add_run(exp.period).italic = True
            
            for hl in exp.highlights:
                doc.add_paragraph(hl, style='List Bullet')

    # Educação
    if resume.education:
        doc.add_heading(SECTION_LABELS["education"], level=1)
        for edu in resume.education:
            p = doc.add_paragraph()
            p.add_run(f"{edu.institution}").bold = True
            p.add_run(f" | {edu.degree}")
            
            p_date = doc.add_paragraph()
            p_date.add_run(edu.period).italic = True
            
            for detail in edu.details:
                doc.add_paragraph(detail, style='List Bullet')

    # Skills
    if resume.skills:
        doc.add_heading(SECTION_LABELS["skills"], level=1)
        for sc in resume.skills:
            p = doc.add_paragraph()
            p.add_run(f"{sc.category}: ").bold = True
            p.add_run(", ".join(sc.items))

    # Projetos
    if resume.projects:
        doc.add_heading(SECTION_LABELS["projects"], level=1)
        for proj in resume.projects:
            p = doc.add_paragraph()
            p.add_run(f"{proj.name}").bold = True
            if proj.technologies:
                p.add_run(f" | {', '.join(proj.technologies)}")
            if proj.url:
                p.add_run(f" ({proj.url})")
            
            doc.add_paragraph(proj.description, style='List Bullet')

    # Idiomas
    if resume.languages:
        doc.add_heading(SECTION_LABELS["languages"], level=1)
        langs = [f"{lang.name}: {lang.level}" for lang in resume.languages]
        doc.add_paragraph(" | ".join(langs))

    # Certificações
    if resume.certifications:
        doc.add_heading(SECTION_LABELS["certifications"], level=1)
        for cert in resume.certifications:
            p = doc.add_paragraph()
            p.add_run(f"{cert.name}").bold = True
            if cert.issuer:
                p.add_run(f" - {cert.issuer}")
            if cert.date:
                p.add_run(f" ({cert.date})")

    # Salva e libera memória
    output_path = OUTPUT_DIR / f"{output_name}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    del doc
    gc.collect()
    
    return output_path
